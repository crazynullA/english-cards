import openpyxl
from docx import Document
from docx.shared import Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_custom_table(doc):
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(6)
        row.height = Cm(3.3)

def fill_surrogate_to_table(table_num):
    currentWord=0
    for row in range(8):
        for col in range(3):
            wordDoc.tables[table_num].rows[row].cells[col].text=str(currentWord)
            currentWord+=1

def replace_surrogate_with_value(table_num,surrogate_name,value):
    for row in range(8):
        for col in range(3):
            if (wordDoc.tables[table_num].rows[row].cells[col].text==str(surrogate_name)):
                wordDoc.tables[table_num].rows[row].cells[col].text=value

def tune_doc_params(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    # ширина листа в сантиметрах
    section.page_width = Cm(21.0)
    # левое поле в миллиметрах
    section.left_margin = Mm(12.7)
    # правое поле в миллиметрах
    section.right_margin = Mm(12.7)
    # верхнее поле в миллиметрах
    section.top_margin = Mm(12.7)
    # нижнее поле в миллиметрах
    section.bottom_margin = Mm(12.7)

path = 'd:\\_Denis\\v.xlsx'
#path_docx = 'd:\\_Denis\\cards.docx'

rus_list=list()
eng_list=list()

wb = openpyxl.load_workbook(path)
ws = wb.active

for rowId in range(1,ws.max_row+1):
    eng_list.append(ws.cell(rowId,1).value)
    rus_list.append(ws.cell(rowId,4).value)

full_lists = ws.max_row//24
notFull_list = ws.max_row%24


wordDoc = Document()
tune_doc_params(wordDoc)

word_index=0
table_counter=0
for full_list_index in range(1,full_lists+1,1):
    print("current list id:",full_list_index)
    add_custom_table(wordDoc)
    add_custom_table(wordDoc)
    fill_surrogate_to_table(table_counter)
    fill_surrogate_to_table(table_counter+1)
    for list_word_index in range(24):
        replace_surrogate_with_value(table_counter,list_word_index,eng_list[word_index])
        replace_surrogate_with_value(table_counter+1,list_word_index,rus_list[word_index])
        word_index+=1    

    #swap words in first and last column
    for row in range(8):
        wordDoc.tables[table_counter+1].rows[row].cells[0].text,wordDoc.tables[table_counter+1].rows[row].cells[2].text = wordDoc.tables[table_counter+1].rows[row].cells[2].text, wordDoc.tables[table_counter+1].rows[row].cells[0].text
    ####################################

    table_counter+=2
       
 
wordDoc.save('d:\\_Crazy\\_workalka\\python\\english-cards\\result.docx')









#wordIdx=0
#for currentList in range(1,full_lists+1,1):
#    for wordIdx in range(currentList*24):
#        print(eng_list[wordIdx])
#    print()
#
#for currentList in range(1,notFull_list+1):
#    print(eng_list[wordIdx+currentList])

wb.close()
